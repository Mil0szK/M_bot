import datetime
import sqlite3

import matplotlib.pyplot as plt
from docx import Document


def calculate_expenses():
    data = get_expenses_data()


def generate_chart(chart_type, data, title, filename):
    plt.style.use("fivethirtyeight")
    if chart_type == 'bar':
        plt.bar(data.keys(), data.values())
    elif chart_type == 'donut':
        fig, ax = plt.subplots()
        ax.pie(data.values(), labels=data.keys(), autopct='%1.1f%%', startangle=90, pctdistance=0.85)
        centre_circle = plt.Circle((0, 0), 0.70, fc='white')
        fig = plt.gcf()
        fig.gca().add_artist(centre_circle)
        plt.axis('equal')
    plt.title(title)
    plt.tight_layout()
    plt.savefig(filename)
    plt.show()


def generate_monthly_bar(month, year):
    # Get the expenses data for the specified month
    monthly_data, _ = get_expenses_for_month(month, year)

    # Get the costs per day
    day_cost = get_costs_per_day(monthly_data)

    generate_chart('bar', day_cost, f"Expenses per day for {month}/{year}", '../monthly_expenses_bar.png')


def generate_yearly_bar(year):
    # Get the expenses data for the specified year
    yearly_data, _ = get_expenses_for_year(year)

    # Get the costs by month
    month_cost = get_costs_per_month(yearly_data)

    # Generate the bar chart
    generate_chart('bar', month_cost, f"Expenses per month for {year}", '../yearly_expenses_bar.png')


def get_expenses_data():
    conn = sqlite3.connect('../data/expenses.db')
    c = conn.cursor()
    c.execute("SELECT * FROM expenses")
    data = c.fetchall()
    conn.close()
    return data


def get_costs_by_category(data):
    # Get the unique categories from the data
    categories = set(row[1] for row in data)

    # Initialize a dictionary to store the costs per category
    cat_cost = {}

    # Calculate the total cost for each category
    for cat in categories:
        cat_cost[cat] = sum(row[2] for row in data if row[1] == cat)

    return cat_cost


def get_costs_per_day(data):
    # Initialize a dictionary to store the costs per day
    day_cost = {day: 0 for day in range(1, 32)}

    # Calculate the total cost for each day
    for row in data:
        day = row[0].day
        day_cost[day] += row[2]

    return day_cost


def get_costs_per_month(data):
    # Initialize a dictionary to store the costs per month
    month_cost = {month: 0 for month in range(1, 13)}

    # Calculate the total cost for each month
    for row in data:
        month = row[0].month
        month_cost[month] += row[2]

    return month_cost


def get_expenses_for_month(month, year):
    # Fetch the data from the database
    data = get_expenses_data()

    # Convert the date strings into datetime objects
    dates = [datetime.datetime.strptime(row[4], '%d.%m.%Y') for row in data]
    categories = [row[1] for row in data]
    amounts = [row[3] for row in data]
    shared = [row[2] for row in data]

    # Filter the data based on the month and year of the datetime objects
    monthly_data = []
    monthly_shared_data = []
    for date, category, amount, is_shared in zip(dates, categories, amounts, shared):
        if date.month == month and date.year == year:
            if is_shared == "yes":
                amount /= 2
                monthly_shared_data.append((date, category, amount))
            monthly_data.append((date, category, amount))

    return monthly_data, monthly_shared_data


def get_expenses_for_year(year):
    # Fetch the data from the database
    data = get_expenses_data()

    # Convert the date strings into datetime objects
    dates = [datetime.datetime.strptime(row[4], '%d.%m.%Y') for row in data]
    categories = [row[1] for row in data]
    amounts = [row[3] for row in data]
    shared = [row[2] for row in data]

    # Filter the data based on the year of the datetime objects
    yearly_data = []
    yearly_shared_data = []
    for date, category, amount, is_shared in zip(dates, categories, amounts, shared):
        if date.year == year:
            if is_shared == "yes":
                amount /= 2
                yearly_shared_data.append((date, category, amount))
            yearly_data.append((date, category, amount))

    return yearly_data, yearly_shared_data


def generate_monthly_donut(month, year):
    # Get the expenses data for the specified month
    monthly_data, _ = get_expenses_for_month(month, year)

    # Get the costs by category
    cat_cost = get_costs_by_category(monthly_data)

    generate_chart('donut', cat_cost, f"Expenses per category for {month}/{year}", '../monthly_expenses_donut.png')


def generate_yearly_donut(year):
    # Get the expenses data for the specified year
    yearly_data, _ = get_expenses_for_year(year)

    # Get the costs by category
    cat_cost = get_costs_by_category(yearly_data)

    # Generate the donut chart
    generate_chart('donut', cat_cost, f"Expenses per category for {year}", '../yearly_expenses_donut.png')


def generate_monthly_shared_donut(month, year):
    # Get the shared expenses data for the specified month
    _, monthly_shared_data = get_expenses_for_month(month, year)

    # Get the costs by category
    cat_cost = get_costs_by_category(monthly_shared_data)

    # Generate the donut chart
    fig, ax = plt.subplots()
    ax.pie(cat_cost.values(), labels=cat_cost.keys(), autopct='%1.1f%%', startangle=90, pctdistance=0.85)
    centre_circle = plt.Circle((0, 0), 0.70, fc='white')
    fig = plt.gcf()
    fig.gca().add_artist(centre_circle)
    plt.axis('equal')
    plt.tight_layout()
    plt.savefig('monthly_shared_expenses_donut.png')
    plt.show()


def generate_yearly_shared_donut(year):
    # Get the shared expenses data for the specified year
    _, yearly_shared_data = get_expenses_for_year(year)

    # Get the costs by category
    cat_cost = get_costs_by_category(yearly_shared_data)

    # Generate the donut chart
    fig, ax = plt.subplots()
    ax.pie(cat_cost.values(), labels=cat_cost.keys(), autopct='%1.1f%%', startangle=90, pctdistance=0.85)
    centre_circle = plt.Circle((0, 0), 0.70, fc='white')
    fig = plt.gcf()
    fig.gca().add_artist(centre_circle)
    plt.axis('equal')
    plt.tight_layout()
    plt.savefig('yearly_shared_expenses_donut.png')
    plt.show()


def generate_monthly_shared_bar(month, year):
    # Get the shared expenses data for the specified month
    _, monthly_shared_data = get_expenses_for_month(month, year)

    # Get the costs per day
    day_cost = get_costs_per_day(monthly_shared_data)

    # Generate the bar chart
    plt.style.use("fivethirtyeight")
    plt.bar(day_cost.keys(), day_cost.values())
    plt.title(f"Shared expenses per day for {month}/{year}")
    plt.tight_layout()
    plt.savefig('monthly_shared_expenses_bar.png')
    plt.show()


def generate_yearly_shared_bar(year):
    # Get the shared expenses data for the specified year
    _, yearly_shared_data = get_expenses_for_year(year)

    # Get the costs per month
    month_cost = get_costs_per_month(yearly_shared_data)

    # Generate the bar chart
    plt.style.use("fivethirtyeight")
    plt.bar(month_cost.keys(), month_cost.values())
    plt.title(f"Shared expenses per month for {year}")
    plt.tight_layout()
    plt.savefig('yearly_shared_expenses_bar.png')
    plt.show()


def generate_yearly_reports(year):
    # Generate the yearly reports
    generate_yearly_bar(year)
    generate_yearly_donut(year)
    generate_yearly_shared_bar(year)
    generate_yearly_shared_donut(year)

    # Create a new Word document
    doc = Document()

    # Add each image to the Word document
    doc.add_picture('yearly_expenses_bar.png')
    doc.add_picture('yearly_expenses_donut.png')
    doc.add_picture('yearly_shared_expenses_bar.png')
    doc.add_picture('yearly_shared_expenses_donut.png')

    # Save the Word document
    doc.save('yearly_reports.docx')

    print("Yearly reports generated and saved in 'yearly_reports.docx'")


def generate_monthly_reports(month, year):
    # Generate the yearly reports
    generate_monthly_bar(month, year)
    generate_monthly_donut(month, year)
    generate_monthly_shared_bar(month, year)
    generate_monthly_shared_donut(month, year)

    # Create a new Word document
    doc = Document()

    # Add each image to the Word document
    doc.add_picture('monthly_expenses_bar.png')
    doc.add_picture('monthly_expenses_donut.png')
    doc.add_picture('monthly_shared_expenses_bar.png')
    doc.add_picture('monthly_shared_expenses_donut.png')

    # Save the Word document
    doc.save('monthly_reports.docx')

    print("Monthly reports generated and saved in 'monthly_reports.docx'")